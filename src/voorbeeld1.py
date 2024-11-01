from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER, WD_BREAK
from helper import set_cell_border, hide_table_borders, set_tokaio_paragraph_style


document = Document()

sections = document.sections
for section in sections:
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

# Create a new header for the first section
header = sections[0].header
run = header.paragraphs[0].add_run()
run.add_picture('images/tokaio.png', width=Cm(7.5))
header.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run.add_break()

hoofd = header.add_table(rows=2, cols=3, width=Cm(19))

rij1 = hoofd.rows[0].cells
run = rij1[0].paragraphs[0].add_run()
run.add_text('Factuur')

run = rij1[2].paragraphs[0].add_run()
run.add_text('Factuur adres')

rij2 = hoofd.rows[1].cells
run = rij2[0].paragraphs[0].add_run()

run.add_text(f'Factuur nr.')
run.add_break(WD_BREAK.LINE)
run.add_text('Factuur datum')
run.add_break(WD_BREAK.LINE)
run.add_text(f'Factuur nr.')
run.add_break(WD_BREAK.LINE)
run.add_text(f'Betalings datum')

run = rij2[2].paragraphs[0].add_run()

run.add_text('Kerkstraat 1')
run.add_break(WD_BREAK.LINE)
run.add_text(f'1234 AB Amsterdam')
run.add_break(WD_BREAK.LINE)
run.add_text(f'Nederland')

hide_table_borders(hoofd)

# Factuur detail
aantal_artikels = 3
document.add_paragraph('Factuur details')
detail_tabel = document.add_table(rows=aantal_artikels + 2, cols=5)

detail_tabel.autofit = False 
detail_tabel.allow_autofit = False

# 19 cm te verdelen over 5 kolommen (3.8)
# Product beschrijving
detail_tabel.columns[0].width = Cm(6.5)
# detail_tabel.rows[0].cells[0].width = Cm(4.5)
detail_tabel.rows[0].cells[0].text = 'Product beschrijving'

# Aantal
detail_tabel.columns[1].width = Cm(2.5)
detail_tabel.rows[0].cells[1].text = 'Aantal'

# Prijs
detail_tabel.columns[2].width = Cm(3)
detail_tabel.rows[0].cells[2].text = 'Prijs'

# BTW
detail_tabel.columns[3].width = Cm(3.5)
detail_tabel.rows[0].cells[3].text = 'BTW (21%)'

# Bedrag
detail_tabel.columns[4].width = Cm(3.5)
detail_tabel.rows[0].cells[4].text = 'Bedrag'

# details
detail_tabel.rows[1].cells[0].text = 'Product 1'
detail_tabel.rows[2].cells[0].text = 'Product 2'
detail_tabel.rows[3].cells[0].text = 'Product 3'

# Aantal
detail_tabel.rows[1].cells[1].text = '1'
detail_tabel.rows[2].cells[1].text = '2'
detail_tabel.rows[3].cells[1].text = '3'

# Prijs
detail_tabel.rows[1].cells[2].text = '€ 100'
detail_tabel.rows[2].cells[2].text = '€ 200'
detail_tabel.rows[3].cells[2].text = '€ 300'

# BTW
detail_tabel.rows[1].cells[3].text = '€ 21'
detail_tabel.rows[2].cells[3].text = '€ 42'
detail_tabel.rows[3].cells[3].text = '€ 63'

# Bedrag
detail_tabel.rows[1].cells[4].text = '€ 121'
detail_tabel.rows[2].cells[4].text = '€ 242'
detail_tabel.rows[3].cells[4].text = '€ 363'

# Totaal
run = detail_tabel.rows[4].cells[3].paragraphs[0].add_run()
run.text = 'Subtotaal'
run.add_break(WD_BREAK.LINE)
run.add_text(f'BTW')
run.add_break(WD_BREAK.LINE)
run.add_text(f'Totaal')

run = detail_tabel.rows[4].cells[4].paragraphs[0].add_run()
run.text = '€ 726'
run.add_break(WD_BREAK.LINE)
run.add_text(f'€ 126')
run.add_break(WD_BREAK.LINE)
run.add_text(f'€ 852')


hide_table_borders(detail_tabel)

trailing_text = document.add_paragraph('Hier komt nog wat tekst onder de tabel')

footer = sections[0].footer
run = footer.paragraphs[0].add_run()
run.add_text('Tokaio BV')
run.add_break()
run.add_text('Kerkstraat 1')
run.add_break()
run.add_text('Nog wa info ')
run.add_break()
run.add_text('En hoplaaa')


document.add_page_break()

document.save('files/invoices/demo.docx')