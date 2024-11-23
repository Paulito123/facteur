from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from enumerations import BorderTemplate, InvoiceTemplate, DocumentType, OfferTemplate
from typing import Dict
import subprocess


class DocHelper:

    def __init__(self):
        pass


    def format_number(self, number: float, decimal_places: int = 2) -> str:
        if decimal_places == 0:
            return f"{number:,.0f}".replace(',','*').replace('.', ',').replace('*','.')
        elif decimal_places == 1:
            return f"{number:,.1f}".replace(',','*').replace('.', ',').replace('*','.')
        elif decimal_places == 2:
            return f"{number:,.2f}".replace(',','*').replace('.', ',').replace('*','.')
    

    def set_cell_border(self, cell, **kwargs):
        """
        Set cell`s border
        Usage:

        set_cell_border(
            cell,
            top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
            bottom={"sz": 12, "color": "#00FF00", "val": "single"},
            start={"sz": 24, "val": "dashed", "shadow": "true"},
            end={"sz": 12, "val": "dashed"},
        )
        """
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        # check for tag existnace, if none found, then create one
        tcBorders = tcPr.first_child_found_in("w:tcBorders")
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)

        # list over all available tags
        for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
            edge_data = kwargs.get(edge)
            if edge_data:
                tag = 'w:{}'.format(edge)

                # check for tag existnace, if none found, then create one
                element = tcBorders.find(qn(tag))
                if element is None:
                    element = OxmlElement(tag)
                    tcBorders.append(element)

                # looks like order of attributes is important
                for key in ["sz", "val", "color", "space", "shadow"]:
                    if key in edge_data:
                        element.set(qn('w:{}'.format(key)), str(edge_data[key]))


    def set_table_border_template(self, table, template: BorderTemplate):
        if template == BorderTemplate.NO_BORDERS:
            # Make table borders invisible
            for row in table.rows:
                for cell in row.cells:
                    self.set_cell_border(
                        cell,
                        top={"sz": 1, "val": "single", "color": "#FFFFFF", "space": "0"},
                        bottom={"sz": 1, "val": "single", "color": "#FFFFFF", "space": "0"},
                        start={"sz": 1, "val": "single", "color": "#FFFFFF", "space": "0"},
                        end={"sz": 1, "val": "single", "color": "#FFFFFF", "space": "0"},
                    )
        elif template == BorderTemplate.DETAIL_1:
            aantal_rijen = len(table.rows)
            counter = 1
            # Set borders for table
            for row in table.rows:
                if counter == 1:
                    for cell in row.cells:
                        self.set_cell_border(
                            cell,
                            top={"sz": 6, "val": "single", "color": "#000000", "space": "0"},
                            bottom={"sz": 6, "val": "single", "color": "#000000", "space": "0"},
                            start={"sz": 1, "val": "single", "color": "#FFFFFF", "space": "0"},
                            end={"sz": 1, "val": "single", "color": "#FFFFFF", "space": "0"},
                        )
                elif counter != aantal_rijen:
                    for cell in row.cells:
                        self.set_cell_border(
                            cell,
                            top={"sz": 1, "val": "single", "color": "#FFFFFF", "space": "0"},
                            bottom={"sz": 1, "val": "single", "color": "#FFFFFF", "space": "0"},
                            start={"sz": 1, "val": "single", "color": "#FFFFFF", "space": "0"},
                            end={"sz": 1, "val": "single", "color": "#FFFFFF", "space": "0"},
                        )
                else:
                    for cell in row.cells:
                        self.set_cell_border(
                            cell,
                            top={"sz": 6, "val": "single", "color": "#000000", "space": "0"},
                            bottom={"sz": 1, "val": "single", "color": "#FFFFFF", "space": "0"},
                            start={"sz": 1, "val": "single", "color": "#FFFFFF", "space": "0"},
                            end={"sz": 1, "val": "single", "color": "#FFFFFF", "space": "0"},
                        )
                counter += 1


    def convert_to_pdf(self, docx_path, pdf_path):
        subprocess.run(['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', '/'.join(pdf_path.split('/')[:-1]), docx_path])


    def create_element(self, name):
        return OxmlElement(name)


    def create_attribute(self, element, name, value):
        element.set(qn(name), value)


    def add_page_number(self, run):
        fldChar1 = self.create_element('w:fldChar')
        self.create_attribute(fldChar1, 'w:fldCharType', 'begin')

        instrText = self.create_element('w:instrText')
        self.create_attribute(instrText, 'xml:space', 'preserve')
        instrText.text = "PAGE"

        fldChar2 = self.create_element('w:fldChar')
        self.create_attribute(fldChar2, 'w:fldCharType', 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)


    def set_styles(self, document: Document) -> None:
        titel_1_stijl = document.styles['Heading 1']
        titel_1_stijl.font.size = Pt(20)
        titel_1_stijl.font.name = 'Liberation Sans'
        titel_1_stijl.font.bold = False
        titel_1_stijl.font.color.rgb = RGBColor(0, 0, 0)
        titel_1_stijl.paragraph_format.space_after = Pt(0)
        titel_1_stijl.paragraph_format.space_before = Pt(0)
        titel_1_stijl.paragraph_format.line_spacing = 0

        titel_2_stijl = document.styles['Heading 2']
        titel_2_stijl.font.size = Pt(14)
        titel_2_stijl.font.bold = False
        titel_2_stijl.font.color.rgb = RGBColor(0, 0, 0)
        titel_2_stijl.paragraph_format.space_after = Pt(12)
        titel_2_stijl.paragraph_format.space_before = Pt(12)
        # titel_2_stijl.paragraph_format.line_spacing = 0

        tabel_hoofding_stijl = document.styles['Heading 3']
        tabel_hoofding_stijl.font.size = Pt(12)
        tabel_hoofding_stijl.font.bold = True
        tekst_stijl = document.styles['Normal']
        tekst_stijl.font.size = Pt(12)
        tekst_stijl.font.name = 'Liberation Sans Narrow'
        footer_stijl = document.styles['Footer']
        footer_stijl.font.name = 'Liberation Sans Narrow'
        footer_stijl.font.size = Pt(10)
        footer_stijl.paragraph_format.space_before = Pt(12)

        # Set the margins for the document
        for section in document.sections:
            section.left_margin = Cm(1.5)
            section.right_margin = Cm(1.5)
        

    def set_header(self, document: Document, data: Dict) -> None:
        # Create a new header for the first section
        document.sections[0].different_first_page_header_footer = True
        header = document.sections[0].first_page_header

        hoofd = header.add_table(rows=2, cols=3, width=Cm(19))

        rij0 = hoofd.rows[0].cells
        titel_cell = rij0[0]
        titel_cell.paragraphs[0].style = document.styles['Heading 1']
        run = titel_cell.paragraphs[0].add_run()
        run.add_text(data["title"])

        bcell = rij0[1]
        ccell = rij0[2]
        image_cell = bcell.merge(ccell)
        image_cell.paragraphs[0].add_run().add_picture(data["path_image"], width=Cm(7.5))
        image_cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        rij1 = hoofd.rows[1].cells
        details_cell = rij1[0]
        run = details_cell.paragraphs[0].add_run()

        if 'period' in data:
            delivery_date_label = "Periode"
            delivery_date = data["period"]
            tab_count = 2
        else:
            delivery_date_label = "Leveringsdatum"
            delivery_date = data["delivery_date"]
            tab_count = 1
        
        run.add_text(f'Factuur nr.')
        run.add_text(f'\t')
        run.add_text(data["invoice_nr"])
        run.add_break(WD_BREAK.LINE)

        run.add_text(f'Factuurdatum')
        run.add_text(f'\t')
        run.add_text(data["invoice_date"])
        run.add_break(WD_BREAK.LINE)

        run.add_text(f'{delivery_date_label}')
        run.add_text(f'\t'*tab_count)
        run.add_text(delivery_date)
        run.add_break(WD_BREAK.LINE)

        run.add_text(f'Vervaldag')
        run.add_text(f'\t')
        run.add_text(data["due_date"])

        run = rij1[2].paragraphs[0].add_run()

        run.add_text(f"{data['debtor_name']}")
        run.add_break(WD_BREAK.LINE)

        if 'debtor_tav'in data:
            run.add_text(f't.a.v. {data["debtor_tav"]}')
            run.add_break(WD_BREAK.LINE)

        run.add_text(f"{data['debtor_street']} {data['debtor_nr']}")
        run.add_break(WD_BREAK.LINE)
        run.add_text(f'{data["debtor_zip"]} {data["debtor_city"]}')
        run.add_break(WD_BREAK.LINE)
        run.add_text(f'{data["debtor_country"]}')

        self.set_table_border_template(hoofd, BorderTemplate.NO_BORDERS)


    def set_body(self, document: Document, data: Dict) -> None:
        # Factuur detail
        aantal_artikels = len(data["items"].keys())

        detail_tabel_titel = document.add_paragraph('Details')
        detail_tabel_titel.style = document.styles['Heading 2']
        detail_tabel = document.add_table(rows=aantal_artikels + 2, cols=6)

        detail_tabel.autofit = False 
        detail_tabel.allow_autofit = False

        # HEADER
        # 19 cm te verdelen over 5 kolommen (3.8) 2.5
        # Product beschrijving
        detail_tabel.columns[0].width = Cm(5.5)
        detail_tabel.rows[0].cells[0].text = 'Product beschrijving'

        # Aantal
        detail_tabel.columns[1].width = Cm(1.5)
        detail_tabel.rows[0].cells[1].text = 'Aantal'

        # Eenheidsprijs
        detail_tabel.columns[2].width = Cm(3)
        detail_tabel.rows[0].cells[2].text = 'Eenheidsprijs'

        # Prijs
        detail_tabel.columns[3].width = Cm(3)
        detail_tabel.rows[0].cells[3].text = 'Bedrag excl. BTW'

        # BTW
        detail_tabel.columns[4].width = Cm(3)
        detail_tabel.rows[0].cells[4].text = 'BTW (21%)'

        # Bedrag
        detail_tabel.columns[5].width = Cm(3)
        detail_tabel.rows[0].cells[5].text = 'Bedrag incl. BTW'

        # DETAILS
        counter = 1
        for item_nr in data["items"].keys():
            detail_tabel.rows[counter].cells[0].text = f'{data["items"][item_nr]["description"]}'
            detail_tabel.rows[counter].cells[1].text = f'{self.format_number(data["items"][item_nr]["qty"], 0)}'
            detail_tabel.rows[counter].cells[2].text = f'{data["symbol"]} {self.format_number(data["items"][item_nr]["unit_amt"])}'
            detail_tabel.rows[counter].cells[3].text = f'{data["symbol"]} {self.format_number(data["items"][item_nr]["base_amt"])}'
            detail_tabel.rows[counter].cells[4].text = f'{data["symbol"]} {self.format_number(data["items"][item_nr]["vat_amt"])}'
            detail_tabel.rows[counter].cells[5].text = f'{data["symbol"]} {self.format_number(data["items"][item_nr]["total_amt"])}'
            counter += 1

        # Totaal
        detail_tabel.rows[aantal_artikels+1].cells[4].paragraphs[0].add_run('Subtotaal')
        detail_tabel.rows[aantal_artikels+1].cells[4].paragraphs[0].add_run().add_break(WD_BREAK.LINE)
        detail_tabel.rows[aantal_artikels+1].cells[4].paragraphs[0].add_run('BTW')
        detail_tabel.rows[aantal_artikels+1].cells[4].paragraphs[0].add_run().add_break(WD_BREAK.LINE)
        detail_tabel.rows[aantal_artikels+1].cells[4].paragraphs[0].add_run('Totaal').bold = True

        detail_tabel.rows[aantal_artikels+1].cells[5].paragraphs[0].add_run(f'{data["symbol"]} {self.format_number(data["invoice_base_amt"])}').bold = False
        detail_tabel.rows[aantal_artikels+1].cells[5].paragraphs[0].add_run().add_break(WD_BREAK.LINE)
        detail_tabel.rows[aantal_artikels+1].cells[5].paragraphs[0].add_run(f'{data["symbol"]} {self.format_number(data["invoice_vat_amt"])}')
        detail_tabel.rows[aantal_artikels+1].cells[5].paragraphs[0].add_run().add_break(WD_BREAK.LINE)
        detail_tabel.rows[aantal_artikels+1].cells[5].paragraphs[0].add_run(f'{data["symbol"]} {self.format_number(data["invoice_total_amt"])}').bold = True

        self.set_table_border_template(detail_tabel, BorderTemplate.DETAIL_1)

        trailing_text = document.add_paragraph()
        run = trailing_text.add_run()
        run.add_break()

        if "payment_date" in data:
            run.add_text(f'Dit factuur is betaald op {data["payment_date"]}.')
        else:
            run.add_text(f'Gelieve het factuurbedrag van {data["symbol"]} {self.format_number(data["invoice_total_amt"])} te betalen voor {data["due_date"]} op rekeningnummer {data["creditor_bank_account"]}.')
        

    def set_footer(self, document: Document, data: Dict) -> None:
        # FOOTERS
        footer = document.sections[0].first_page_footer
        voet_1_tabel = footer.add_table(rows=1, cols=3, width=Cm(19))

        rij0 = voet_1_tabel.rows[0].cells
        voet_1_cell_1 = rij0[0]
        run = voet_1_cell_1.paragraphs[0].add_run()
        run.add_text(f'{data["creditor_name"]}')
        run.add_break()
        run.add_text(f"{data['creditor_street']} {data['creditor_nr']}")
        run.add_break()
        run.add_text(f'{data["creditor_zip"]} {data["creditor_city"]} {data["creditor_country"]}')
        voet_1_cell_2 = rij0[1]
        run = voet_1_cell_2.paragraphs[0].add_run()
        run.add_text(f'rpr {data["creditor_rpr"]}')
        run.add_break()
        run.add_text(f'btw {data["creditor_vat"]}')
        run.add_break()
        run.add_text(f'rek {data["creditor_bank_account"]}')
        voet_1_cell_3 = rij0[2]
        run = voet_1_cell_3.paragraphs[0].add_run()
        run.add_text(f'tel {data["creditor_phone"]}')
        run.add_break()
        run.add_text(f'{data["creditor_email"]}')
        run.add_break()
        run.add_text('Pagina ')
        self.add_page_number(run)

        other_footer = document.sections[0].footer
        voet_2_tabel = other_footer.add_table(rows=1, cols=3, width=Cm(19))

        rij0 = voet_2_tabel.rows[0].cells
        voet_2_cell_1 = rij0[0]
        run = voet_2_cell_1.paragraphs[0].add_run()
        run.add_text(f'{data["creditor_name"]}')
        run.add_break()
        run.add_text(f"{data['creditor_street']} {data['creditor_nr']}")
        run.add_break()
        run.add_text(f'{data["creditor_zip"]} {data["creditor_city"]} {data["creditor_country"]}')
        voet_2_cell_2 = rij0[1]
        run = voet_2_cell_2.paragraphs[0].add_run()
        run.add_text(f'rpr {data["creditor_rpr"]}')
        run.add_break()
        run.add_text(f'btw {data["creditor_vat"]}')
        run.add_break()
        run.add_text(f'rek {data["creditor_bank_account"]}')
        voet_2_cell_3 = rij0[2]
        run = voet_2_cell_3.paragraphs[0].add_run()
        run.add_text(f'tel {data["creditor_phone"]}')
        run.add_break()
        run.add_text(f'{data["creditor_email"]}')
        run.add_break()
        run.add_text('Pagina ')
        self.add_page_number(run)

        voet_1_cell_1.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        voet_1_cell_1.paragraphs[0].style = document.styles['Footer']
        voet_1_cell_2.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        voet_1_cell_2.paragraphs[0].style = document.styles['Footer']
        voet_1_cell_3.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        voet_1_cell_3.paragraphs[0].style = document.styles['Footer']
        voet_2_cell_1.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        voet_2_cell_1.paragraphs[0].style = document.styles['Footer']
        voet_2_cell_2.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        voet_2_cell_2.paragraphs[0].style = document.styles['Footer']
        voet_2_cell_3.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        voet_2_cell_3.paragraphs[0].style = document.styles['Footer']
